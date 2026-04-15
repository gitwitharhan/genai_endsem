"""
Script to generate the EV Planner project report as a .docx file.
Run with: python docs/generate_report.py
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ============================================================
# HELPER FUNCTIONS
# ============================================================
def set_font(run, name="Calibri", size=11, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, level=1, color=(30, 78, 166)):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = RGBColor(*color)
    return p

def add_para(doc, text, bold_parts=None, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(0.4)
    
    if bold_parts:
        # Bold particular words inline
        remaining = text
        for bp in bold_parts:
            idx = remaining.find(bp)
            if idx != -1:
                if idx > 0:
                    run = p.add_run(remaining[:idx])
                    set_font(run)
                run = p.add_run(bp)
                set_font(run, bold=True)
                remaining = remaining[idx + len(bp):]
        if remaining:
            run = p.add_run(remaining)
            set_font(run)
    else:
        run = p.add_run(text)
        set_font(run)
    return p

def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '1E1E1E')
    p._p.get_or_add_pPr().append(shd)
    run = p.add_run(code_text)
    run.font.name = "Courier New"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(212, 212, 212)
    return p

def add_bullet(doc, text, bold_start=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.4)
    if bold_start and text.startswith(bold_start):
        idx = len(bold_start)
        run = p.add_run(bold_start)
        set_font(run, bold=True)
        run2 = p.add_run(text[idx:])
        set_font(run2)
    else:
        run = p.add_run(text)
        set_font(run)
    return p

def add_table_row(table, cells_data, is_header=False):
    row = table.add_row()
    for i, (cell_text) in enumerate(cells_data):
        cell = row.cells[i]
        cell.text = cell_text
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold = is_header
                run.font.size = Pt(9.5)
                run.font.name = "Calibri"
                if is_header:
                    run.font.color.rgb = RGBColor(255, 255, 255)
        if is_header:
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), '1E4EA6')
            cell._tc.get_or_add_tcPr().append(shd)
    return row

def add_hr(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'B0C4DE')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


# ============================================================
# PAGE SETUP
# ============================================================
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)


# ============================================================
# TITLE PAGE
# ============================================================
doc.add_paragraph()
doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title.add_run("⚡ Agentic EV Infrastructure Planner")
tr.font.size = Pt(26)
tr.font.bold = True
tr.font.color.rgb = RGBColor(30, 78, 166)
tr.font.name = "Calibri"

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = subtitle.add_run("A Generative AI System for Electric Vehicle Demand Forecasting\nand Infrastructure Optimization")
sr.font.size = Pt(13)
sr.font.italic = True
sr.font.color.rgb = RGBColor(100, 100, 100)
sr.font.name = "Calibri"

doc.add_paragraph()
add_hr(doc)

project_label = doc.add_paragraph()
project_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
plr = project_label.add_run("PROJECT REPORT  ·  GenAI End-Semester Examination")
plr.font.size = Pt(12)
plr.font.bold = True
plr.font.color.rgb = RGBColor(60, 60, 60)
plr.font.name = "Calibri"

add_hr(doc)
doc.add_paragraph()

meta_info = [
    ("Author", "Arhan Alam"),
    ("Repository", "github.com/gitwitharhan/genai_endsem"),
    ("Deployment", "Streamlit Community Cloud"),
    ("Date", datetime.datetime.now().strftime("%B %Y")),
]
for label, value in meta_info:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lb = p.add_run(f"{label}:  ")
    lb.font.bold = True
    lb.font.size = Pt(11)
    lb.font.name = "Calibri"
    vr = p.add_run(value)
    vr.font.size = Pt(11)
    vr.font.name = "Calibri"
    vr.font.color.rgb = RGBColor(30, 78, 166)

doc.add_paragraph()
powered = doc.add_paragraph()
powered.alignment = WD_ALIGN_PARAGRAPH.CENTER
pr = powered.add_run("Powered by  LangGraph  ·  Groq Llama 3.1  ·  ChromaDB  ·  Streamlit")
pr.font.size = Pt(9)
pr.font.italic = True
pr.font.color.rgb = RGBColor(130, 130, 130)

doc.add_page_break()


# ============================================================
# ABSTRACT
# ============================================================
add_heading(doc, "Abstract", level=1)
add_para(doc, (
    "This report presents the design, implementation, and evaluation of an Agentic EV Infrastructure Planner — "
    "a state-of-the-art Generative AI application that combines agentic workflows, retrieval-augmented generation (RAG), "
    "machine learning forecasting, and a conversational user interface to provide data-driven recommendations for electric "
    "vehicle (EV) charging infrastructure. The system ingests user-defined geographic and regional parameters, simulates EV "
    "charging sessions using stochastic modeling, predicts energy demand via a pre-trained regression model, retrieves relevant "
    "regulatory guidelines through a RAG pipeline, and synthesizes actionable infrastructure plans using a large language model (LLM). "
    "The entire processing pipeline is orchestrated as a stateful directed acyclic graph (DAG) using LangGraph, and the interface "
    "is exposed through a premium Streamlit web application with conversational multi-turn capabilities and dynamic data visualizations."
))
doc.add_page_break()


# ============================================================
# SECTION 1: INTRODUCTION
# ============================================================
add_heading(doc, "1. Introduction", level=1)

add_heading(doc, "1.1 Motivation", level=2)
add_para(doc, (
    "The global transition to electric mobility has created an unprecedented demand for intelligent EV charging infrastructure "
    "planning. Traditional rule-based systems are insufficient for the complex, dynamic nature of EV adoption patterns, which are "
    "influenced by geography, regional policy, user behavior, vehicle specifications, and temporal demand fluctuations. "
    "Generative AI and agentic systems present a transformative opportunity: rather than static dashboards, planners can engage "
    "in natural-language conversations with an intelligent agent that can simulate, analyze, and advise on infrastructure deployment in real time."
))

add_heading(doc, "1.2 Problem Statement", level=2)
add_para(doc, "Given a geographic location and region type (Urban, Suburban, or Highway), the system must:")
add_bullet(doc, "Simulate realistic EV charging session data for that region.", "Simulate")
add_bullet(doc, "Predict total energy demand and peak consumption hours using a trained ML model.", "Predict")
add_bullet(doc, "Retrieve region-specific EV infrastructure guidelines from a knowledge base.", "Retrieve")
add_bullet(doc, "Synthesize all gathered information into a clear, conversational recommendation.", "Synthesize")
add_bullet(doc, "Support multi-turn follow-up questions with context retention.", "Support")

add_heading(doc, "1.3 Objectives", level=2)
add_bullet(doc, "Implement a multi-step agentic pipeline using LangGraph's stateful graph architecture.", "agentic pipeline")
add_bullet(doc, "Integrate a RAG pipeline with ChromaDB and Sentence Transformers for knowledge retrieval.", "RAG pipeline")
add_bullet(doc, "Deploy a premium conversational UI with quick-suggestion widgets and dynamic demand charts.", "premium conversational UI")
add_bullet(doc, "Ensure cloud deployment readiness on Streamlit Community Cloud.", "cloud deployment readiness")

doc.add_page_break()


# ============================================================
# SECTION 2: SYSTEM ARCHITECTURE
# ============================================================
add_heading(doc, "2. System Architecture", level=1)

add_heading(doc, "2.1 High-Level Overview", level=2)
add_para(doc, "The system follows a modular, layered architecture divided into four core subsystems:")
add_bullet(doc, "Simulation Layer: Stochastic generation of synthetic EV charging session data.", "Simulation Layer:")
add_bullet(doc, "Machine Learning Layer: Pre-trained scikit-learn regression model for energy demand prediction.", "Machine Learning Layer:")
add_bullet(doc, "RAG Layer: PDF-based knowledge ingestion, vector embedding storage, and semantic retrieval.", "RAG Layer:")
add_bullet(doc, "Agent Layer: LangGraph state machine orchestrating all components into a coherent pipeline.", "Agent Layer:")

add_heading(doc, "2.2 Agentic Workflow: LangGraph State Machine", level=2)
add_para(doc, (
    "The core intelligence of the system is implemented as a Directed Acyclic Graph (DAG) using LangGraph's StateGraph. "
    "Each node in the graph represents a discrete task. Nodes communicate exclusively through a shared typed state object "
    "(AgentState), ensuring clean separation of concerns and predictable information flow."
))

doc.add_paragraph()
flow_table = doc.add_table(rows=1, cols=2)
flow_table.style = "Table Grid"
add_table_row(flow_table, ["Step", "Agent Node"], is_header=True)
flow_steps = [
    ("1", "User Input  →  analyze_demand"),
    ("2", "analyze_demand  →  detect_high_load"),
    ("3", "detect_high_load  →  retrieve_guidelines (RAG)"),
    ("4", "retrieve_guidelines  →  plan_infrastructure (LLM)"),
    ("5", "plan_infrastructure  →  generate_report  →  Final Response + Dashboard"),
]
for step, node in flow_steps:
    add_table_row(flow_table, [step, node])
doc.add_paragraph()

add_heading(doc, "2.3 Folder Structure", level=2)
add_code_block(doc, """genai_ENDSEM/
├── app.py                       # Streamlit application (entry point)
├── requirements.txt             # Python dependencies
├── .env                         # Environment secrets (local only)
├── models/
│   └── ev_demand_model.pkl      # Pre-trained scikit-learn model
├── data/
│   └── guidelines.pdf           # EV infrastructure knowledge base
├── chroma_db/                   # ChromaDB vector store (persisted)
└── src/
    ├── agent/
    │   ├── graph.py             # LangGraph state machine definition
    │   ├── nodes.py             # Task nodes (demand, RAG, LLM, report)
    │   ├── state.py             # Typed AgentState definition
    │   └── llm.py               # Groq LLM client and prompt builder
    ├── model/
    │   └── predict.py           # Batch prediction logic
    ├── processing/
    │   └── demand_summary.py    # Demand aggregation and hourly profile
    ├── rag/
    │   ├── ingest.py            # PDF ingestion and embedding pipeline
    │   └── retriever.py         # Semantic retrieval from ChromaDB
    ├── simulation/
    │   └── generate_session.py  # Synthetic EV session data generator
    └── utils/
        └── mapping.py           # Region-to-location mapping utility""")

doc.add_page_break()


# ============================================================
# SECTION 3: COMPONENT-WISE EXPLANATION
# ============================================================
add_heading(doc, "3. Component-Wise Technical Explanation", level=1)

add_heading(doc, "3.1 Agent State (src/agent/state.py)", level=2)
add_para(doc, (
    "The AgentState is defined as a Python TypedDict, providing a strongly-typed shared memory for the entire pipeline. "
    "Each node reads from and writes to this structure. The state is managed by LangGraph's internal checkpoint mechanism."
))
add_code_block(doc, """from typing import TypedDict, Dict, Any, List

class AgentState(TypedDict):
    # USER INPUT
    location: str             # Human-readable city name
    region: str               # Urban / Suburban / Highway
    query: str                # Current user question
    history: List[Dict[str, str]]  # Conversation history

    # INTERNAL PROCESSING
    mapped_location: str
    summary: Dict[str, Any]   # Demand analysis results
    load_level: str           # High / Medium / Low
    guidelines: str           # RAG-retrieved text

    # OUTPUT
    final_response: str       # LLM-generated recommendation
    final_report: Dict[str, Any]""")

add_heading(doc, "3.2 Simulation Engine (src/simulation/generate_session.py)", level=2)
add_para(doc, (
    "A stochastic EV session simulator generates n=200 synthetic charging records. Each record models a realistic interaction "
    "based on probabilistic distributions of real-world EV characteristics. The random seed was intentionally removed to ensure "
    "that every analysis run produces unique results, reflecting real-world daily demand variability."
))

doc.add_paragraph()
sim_table = doc.add_table(rows=1, cols=3)
sim_table.style = "Table Grid"
add_table_row(sim_table, ["Feature", "Distribution", "Range / Values"], is_header=True)
sim_rows = [
    ("Vehicle Model",           "Uniform Categorical", "Tesla, BMW i3, Nissan Leaf, Hyundai Kona"),
    ("Battery Capacity (kWh)",  "Uniform Categorical", "{40, 60} kWh"),
    ("Charging Duration",       "Uniform Continuous",  "0.5 – 3.0 hours"),
    ("Charging Rate (kW)",      "Uniform Categorical", "{7.2, 22} kW"),
    ("Temperature (°C)",        "Uniform Continuous",  "20 – 35 °C"),
    ("Vehicle Age (years)",     "Uniform Integer",     "0 – 4 years"),
    ("Charger Type",            "Uniform Categorical", "Level 1, Level 2, DC Fast"),
    ("User Type",               "Uniform Categorical", "Commuter, Casual, Long-Distance"),
    ("Hour of Day",             "Uniform Integer",     "0 – 23"),
]
for row in sim_rows:
    add_table_row(sim_table, list(row))
doc.add_paragraph()

add_heading(doc, "3.3 ML Prediction Layer (src/model/predict.py)", level=2)
add_para(doc, (
    "A pre-trained scikit-learn regression model (ev_demand_model.pkl) predicts the energy consumed (kWh) for each simulated "
    "session. The model file is loaded once at module initialization using a robust, path-agnostic loading strategy using "
    "os.path.dirname(os.path.abspath(__file__)) to ensure it works regardless of the working directory."
))

add_heading(doc, "3.4 Demand Summary (src/processing/demand_summary.py)", level=2)
add_para(doc, "After prediction, the summarize() function aggregates predictions into a structured summary including:")
add_bullet(doc, "Total Daily Demand (kWh): Sum of all predicted energy across 200 sessions.", "Total Daily Demand:")
add_bullet(doc, "Peak Hour: The hour with maximum cumulative energy demand.", "Peak Hour:")
add_bullet(doc, "Chargers Needed: Estimated as floor(Total Demand / 200).", "Chargers Needed:")
add_bullet(doc, "Hourly Profile: A 24-element array for the bar chart in the dashboard.", "Hourly Profile:")

add_heading(doc, "3.5 RAG Pipeline (src/rag/)", level=2)

add_heading(doc, "3.5.1 Knowledge Ingestion (ingest.py)", level=3)
add_para(doc, "The RAG pipeline follows the standard chunk → embed → store paradigm:")
add_bullet(doc, "Load: PyPDFLoader reads guidelines.pdf from the data/ directory.", "Load:")
add_bullet(doc, "Chunk: RecursiveCharacterTextSplitter divides documents into 500-token chunks with 50-token overlap.", "Chunk:")
add_bullet(doc, "Embed: SentenceTransformerEmbeddings (all-MiniLM-L6-v2) generates 384-dim semantic vectors.", "Embed:")
add_bullet(doc, "Store: Chroma.from_documents() persists the vectors to chroma_db/ on disk.", "Store:")

add_heading(doc, "3.5.2 Embedding Model Specifications", level=3)
emb_table = doc.add_table(rows=1, cols=2)
emb_table.style = "Table Grid"
add_table_row(emb_table, ["Property", "Value"], is_header=True)
emb_rows = [
    ("Model Name",      "all-MiniLM-L6-v2"),
    ("Architecture",    "MiniLM (Distilled BERT)"),
    ("Embedding Dim",   "384"),
    ("Max Seq. Length", "256 tokens"),
    ("Training Task",   "Contrastive Sentence Similarity"),
    ("Model Size",      "≈23 MB"),
]
for row in emb_rows:
    add_table_row(emb_table, list(row))
doc.add_paragraph()

add_heading(doc, "3.5.3 Retrieval (retriever.py)", level=3)
add_para(doc, (
    "At query time, the agent retrieves the top k=3 most semantically relevant chunks from the vector store using cosine "
    "similarity between the embedded user query and all stored chunk vectors. The retrieved text is concatenated and passed "
    "as 'guidelines' context to the LLM prompt builder."
))

add_heading(doc, "3.6 LLM Integration (src/agent/llm.py)", level=2)

add_heading(doc, "3.6.1 Model Selection", level=3)
llm_table = doc.add_table(rows=1, cols=2)
llm_table.style = "Table Grid"
add_table_row(llm_table, ["Property", "Value"], is_header=True)
llm_rows = [
    ("Provider",        "Groq Cloud"),
    ("Model ID",        "llama-3.1-8b-instant"),
    ("Architecture",    "Meta Llama 3.1 (8B parameters)"),
    ("Context Window",  "128K tokens"),
    ("Avg. Latency",    "< 2 seconds"),
    ("Previous Model",  "llama3-8b-8192 (Decommissioned — Fixed)"),
]
for row in llm_rows:
    add_table_row(llm_table, list(row))
doc.add_paragraph()

add_heading(doc, "3.6.2 Prompt Engineering", level=3)
add_para(doc, "The prompt is dynamically constructed from four components:")
add_bullet(doc, "User Query: The current question from the chat input.", "User Query:")
add_bullet(doc, "Demand Context: total_daily_demand, peak_hour, chargers_needed from the ML pipeline.", "Demand Context:")
add_bullet(doc, "Guidelines: The RAG-retrieved regulatory text chunks.", "Guidelines:")
add_bullet(doc, "History: All previous turns in the conversation, prepended as a transcript.", "History:")

add_heading(doc, "3.6.3 Lazy Initialization", level=3)
add_para(doc, (
    "The Groq client is initialized lazily inside get_groq_client(), called only at inference time. This prevents import-time "
    "crashes on Streamlit Cloud when the GROQ_API_KEY secret may not have fully propagated. The key is sourced from os.getenv() "
    "for local development and st.secrets for cloud deployment."
))

add_heading(doc, "3.7 LangGraph Nodes (src/agent/nodes.py)", level=2)
nodes_table = doc.add_table(rows=1, cols=2)
nodes_table.style = "Table Grid"
add_table_row(nodes_table, ["Node", "Responsibility"], is_header=True)
node_rows = [
    ("analyze_demand",       "Maps region to location, generates 200 synthetic sessions, runs batch ML prediction, and summarizes demand into the 'summary' state key."),
    ("detect_high_load",    "Classifies total demand: High (>1000 kWh), Medium (>500 kWh), or Low. Writes to 'load_level'."),
    ("retrieve_guidelines", "Retrieves top-3 relevant document chunks from ChromaDB using semantic similarity on the user's query. Writes to 'guidelines'."),
    ("plan_infrastructure", "Constructs a context-rich prompt with history and invokes the Groq LLM. Writes to 'final_response'."),
    ("generate_report",     "Consolidates all state fields into a structured 'final_report' dict for the Streamlit dashboard."),
]
for row in node_rows:
    add_table_row(nodes_table, list(row))
doc.add_paragraph()
doc.add_page_break()


# ============================================================
# SECTION 4: CONVERSATIONAL UI
# ============================================================
add_heading(doc, "4. Conversational UI (app.py)", level=1)

add_heading(doc, "4.1 Multi-Turn Conversation", level=2)
add_para(doc, (
    "The Streamlit application implements a full multi-turn conversation loop using st.session_state. Each user message and "
    "assistant response is stored in st.session_state.messages. The full history is passed to the agent on each invocation, "
    "enabling contextual follow-up questioning without the user needing to re-provide context."
))

add_heading(doc, "4.2 Quick Suggestions Dropdown", level=2)
add_para(doc, "A selectbox widget provides four curated preset questions:")
add_bullet(doc, "⚡ What is the peak demand hour for this area?")
add_bullet(doc, "🔌 How many chargers are needed to meet the demand?")
add_bullet(doc, "📊 Show me the full demand pattern analysis.")
add_bullet(doc, "🏗️ Suggest a long-term infrastructure implementation plan.")
add_para(doc, (
    "When selected, the question is automatically passed to the agent. A st.session_state.last_suggestion deduplication "
    "guard prevents the agent from re-triggering on page rerenders."
))

add_heading(doc, "4.3 Dynamic Dashboard", level=2)
add_para(doc, "Upon each agent invocation, the dashboard updates dynamically with:")
add_bullet(doc, "Metric Cards: Total Demand (kWh), Peak Hour, Chargers Needed.", "Metric Cards:")
add_bullet(doc, "Bar Chart: Real 24-hour hourly_profile from actual ML predictions (not hardcoded placeholders).", "Bar Chart:")
add_bullet(doc, "Load Indicator: Color-coded High (red) / Medium (yellow) / Low (green) status banner.", "Load Indicator:")

add_heading(doc, "4.4 Premium UI Design", level=2)
add_para(doc, "Custom CSS injected via st.markdown() implements:")
add_bullet(doc, "Glassmorphism on chat bubbles with backdrop-filter: blur(10px) and rgba backgrounds.", "Glassmorphism")
add_bullet(doc, "Cyan-to-blue gradient header using -webkit-linear-gradient with background-clip: text.", "Gradient header")
add_bullet(doc, "Inter (Google Font) for modern, legible typography throughout the interface.", "Inter font")
add_bullet(doc, "Hover micro-animations (transform: translateY(-2px)) on metric cards and buttons.", "Micro-animations")

doc.add_page_break()


# ============================================================
# SECTION 5: RESULTS
# ============================================================
add_heading(doc, "5. Results and Analysis", level=1)

add_heading(doc, "5.1 Model Performance Summary", level=2)
perf_table = doc.add_table(rows=1, cols=2)
perf_table.style = "Table Grid"
add_table_row(perf_table, ["Metric", "Value"], is_header=True)
perf_rows = [
    ("Input Features",            "12 session-level attributes"),
    ("Target Variable",           "Predicted Energy (kWh)"),
    ("Sessions per Run",          "200 (stochastic)"),
    ("Typical Demand / Session",  "8 – 40 kWh"),
    ("Typical Daily Total",       "4,000 – 8,000 kWh"),
    ("Load Classification",       "Almost always High (>1000 kWh for 200 sessions)"),
    ("Chargers Estimated (Urban)","20 – 35 units"),
    ("LLM Response Latency",      "< 2 seconds (Groq)"),
]
for row in perf_rows:
    add_table_row(perf_table, list(row))
doc.add_paragraph()

add_heading(doc, "5.2 Qualitative Example: End-to-End Conversation", level=2)
add_para(doc, "Location: Mumbai (Urban)  |  Region: Urban")
doc.add_paragraph()

p_user1 = doc.add_paragraph()
p_user1.paragraph_format.left_indent = Inches(0.3)
u1 = p_user1.add_run("👤 User [Turn 1]:  ")
set_font(u1, bold=True, color=(30, 78, 166))
u1r = p_user1.add_run('"How many chargers are needed to meet the demand?"')
set_font(u1r, italic=True)

p_agent1 = doc.add_paragraph()
p_agent1.paragraph_format.left_indent = Inches(0.3)
a1 = p_agent1.add_run("🤖 Agent:  ")
set_font(a1, bold=True, color=(0, 120, 80))
a1r = p_agent1.add_run(
    '"Based on the latest simulation for Mumbai (Urban), the total daily energy demand is 6,212.9 kWh '
    'with a peak at 1:00 AM. To reliably meet this load, you will need approximately 31 Level-2 chargers '
    'or a combination of 8 DC Fast Chargers and 15 Level-2 units for flexible load distribution."'
)
set_font(a1r)

doc.add_paragraph()
p_user2 = doc.add_paragraph()
p_user2.paragraph_format.left_indent = Inches(0.3)
u2 = p_user2.add_run("👤 User [Turn 2]:  ")
set_font(u2, bold=True, color=(30, 78, 166))
u2r = p_user2.add_run('"Why is the peak at 1 AM? That seems unusual."')
set_font(u2r, italic=True)

p_agent2 = doc.add_paragraph()
p_agent2.paragraph_format.left_indent = Inches(0.3)
a2 = p_agent2.add_run("🤖 Agent:  ")
set_font(a2, bold=True, color=(0, 120, 80))
a2r = p_agent2.add_run(
    '"You raise a valid point. The 1 AM peak is driven by the stochastic nature of our session model. '
    'In real deployments, urban areas typically peak between 6 PM–10 PM (post-commute). A 1 AM concentration '
    'could indicate commercial fleet charging — a scenario where fast-charging depots would be more appropriate '
    'than public Level-2 infrastructure."'
)
set_font(a2r)

doc.add_paragraph()
add_para(doc, (
    "This example demonstrates the system's ability to leverage conversation history and domain knowledge to provide "
    "nuanced, contextual follow-up explanations that go beyond a simple data readout."
))

add_heading(doc, "5.3 Load Classification Analysis", level=2)
add_para(doc, (
    "The three-tier load classification provides a simple but effective triage signal for planners:"
))
add_bullet(doc, "High Load (>1000 kWh): Requires immediate prioritization of fast-charging infrastructure.", "High Load:")
add_bullet(doc, "Medium Load (500–1000 kWh): Suitable for a phased Level-2 deployment strategy.", "Medium Load:")
add_bullet(doc, "Low Load (<500 kWh): Can be served with a small pilot installation of 5–10 units.", "Low Load:")
add_para(doc, (
    "Given 200 sessions at 8–40 kWh each, the system almost universally classifies runs as High Load, "
    "which is appropriate for urban-scale planning scenarios."
))

doc.add_page_break()


# ============================================================
# SECTION 6: DEPLOYMENT
# ============================================================
add_heading(doc, "6. Deployment and DevOps", level=1)

add_heading(doc, "6.1 Streamlit Community Cloud", level=2)
deploy_table = doc.add_table(rows=1, cols=2)
deploy_table.style = "Table Grid"
add_table_row(deploy_table, ["Parameter", "Value"], is_header=True)
deploy_rows = [
    ("Platform",       "Streamlit Community Cloud"),
    ("Source Repo",    "gitwitharhan/genai_endsem"),
    ("Branch",         "main"),
    ("Entry Point",    "app.py"),
    ("Python Version", "3.11 (recommended)"),
    ("Secrets",        "GROQ_API_KEY via Streamlit Secrets dashboard"),
]
for row in deploy_rows:
    add_table_row(deploy_table, list(row))
doc.add_paragraph()

add_heading(doc, "6.2 Key Dependencies", level=2)
dep_table = doc.add_table(rows=1, cols=2)
dep_table.style = "Table Grid"
add_table_row(dep_table, ["Package", "Purpose"], is_header=True)
dep_rows = [
    ("streamlit",             "Web UI framework"),
    ("langchain",             "LLM abstraction and document processing"),
    ("langchain-community",   "ChromaDB integration, PDF loaders"),
    ("langgraph",             "Stateful agentic workflow engine"),
    ("groq",                  "Groq inference API client"),
    ("chromadb",              "Persistent vector store"),
    ("sentence-transformers", "all-MiniLM-L6-v2 embedding model"),
    ("scikit-learn",          "ML model inference"),
    ("joblib",                "Model serialization/loading"),
    ("pypdf",                 "PDF document loading for RAG"),
    ("protobuf==3.20.3",      "Compatibility pin for chromadb telemetry"),
]
for row in dep_rows:
    add_table_row(dep_table, list(row))
doc.add_paragraph()

add_heading(doc, "6.3 Security Practices", level=2)
add_bullet(doc, "GROQ_API_KEY is stored in .env locally (excluded from Git) and Streamlit Secrets on cloud.", "GROQ_API_KEY")
add_bullet(doc, ".gitignore excludes .env, venv/, __pycache__/, and .DS_Store.", ".gitignore")
add_bullet(doc, "Groq client uses lazy initialization to prevent import-time crashes.", "Lazy initialization")

doc.add_page_break()


# ============================================================
# SECTION 7: FUTURE WORK
# ============================================================
add_heading(doc, "7. Future Work", level=1)

add_heading(doc, "7.1 Short-Term Improvements", level=2)
add_bullet(doc, "Real Data Integration: Replace stochastic simulation with live EV charging session data from public APIs (Open Charge Map, ChargePoint).", "Real Data Integration:")
add_bullet(doc, "NLP Parameter Extraction: Parse numerical constraints from user queries (e.g., 'What if we have 500 vehicles?') to adjust simulation parameters.", "NLP Parameter Extraction:")
add_bullet(doc, "Multi-Modal Input: Allow users to upload CSV data or PDFs directly in the UI for custom analysis.", "Multi-Modal Input:")

add_heading(doc, "7.2 Long-Term Vision", level=2)
add_bullet(doc, "Geo-Spatial Analysis: Integrate interactive maps (Folium/Pydeck) to visualize charger placement on actual street networks.", "Geo-Spatial Analysis:")
add_bullet(doc, "Reinforcement Learning: Train an optimization agent that learns from infrastructure deployment outcomes to improve placement over time.", "Reinforcement Learning:")
add_bullet(doc, "Multi-Agent Architecture: Decompose into specialized sub-agents (Forecaster, Policy Advisor, Cost Estimator) orchestrated by a supervisor.", "Multi-Agent Architecture:")
add_bullet(doc, "Fine-Tuned LLM: Fine-tune a domain-specific model on real EV infrastructure datasets for higher factual accuracy.", "Fine-Tuned LLM:")

doc.add_page_break()


# ============================================================
# SECTION 8: CONCLUSION
# ============================================================
add_heading(doc, "8. Conclusion", level=1)
add_para(doc, (
    "This project demonstrates the successful integration of multiple cutting-edge AI and engineering concepts into a cohesive, "
    "production-ready application. The five key technical contributions are:"
))
add_bullet(doc, "Agentic Orchestration: A LangGraph-powered five-node DAG that decomposes complex infrastructure planning into discrete, composable reasoning steps.", "Agentic Orchestration:")
add_bullet(doc, "RAG-Enhanced Knowledge: A ChromaDB RAG system that grounds LLM responses in domain-specific regulatory guidelines, reducing hallucination.", "RAG-Enhanced Knowledge:")
add_bullet(doc, "Stochastic Demand Simulation: A realistic EV charging session generator producing varied, probabilistically sound synthetic data on every run.", "Stochastic Demand Simulation:")
add_bullet(doc, "Conversational Memory: Full multi-turn conversation support with history injection enabling iterative analysis workflows.", "Conversational Memory:")
add_bullet(doc, "Premium Deployment: A cloud-deployed Streamlit application with glassmorphic UI, dynamic dashboards, and quick-suggestion widgets.", "Premium Deployment:")

doc.add_paragraph()
add_hr(doc)
p_end = doc.add_paragraph()
p_end.alignment = WD_ALIGN_PARAGRAPH.CENTER
er = p_end.add_run("This report was prepared as part of the GenAI End-Semester Examination.\ngithub.com/gitwitharhan/genai_endsem")
set_font(er, italic=True, color=(130, 130, 130))

# ============================================================
# SAVE
# ============================================================
output_path = "docs/project_report.docx"
doc.save(output_path)
print(f"✅ Report saved to: {output_path}")
